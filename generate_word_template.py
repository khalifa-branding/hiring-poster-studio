import docx
from docx.shared import Inches, Pt, RGBColor, Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

OFFICE_DATA = {
    "bangalore": {
        "filename": "Trescon_Executive_Letterhead_Bangalore.docx",
        "blank_filename": "Trescon_Letterhead_Bangalore_Blank.docx",
        "entity": "Trescon Global Business Solutions Pvt Ltd.",
        "address": "1st floor, Prom’S Complex, 3h, 7th C Main Rd, 3rd Block Koramangala, Bengaluru, Karnataka – 560034",
        "extra": "",
        "cin": "CIN: U74900KA2016PTC086221",
        "email": "info@tresconglobal.com",
        "web": "tresconglobal.com"
    },
    "manipal": {
        "filename": "Trescon_Executive_Letterhead_Manipal.docx",
        "blank_filename": "Trescon_Letterhead_Manipal_Blank.docx",
        "entity": "Trescon Global Business Solutions Pvt Ltd.",
        "address": "H (23), 5th Floor, Pragathi Business District #412, above Reliance Trends, Laxmindra Nagar, Manipal, Udupi, Karnataka – 576104",
        "extra": "",
        "cin": "CIN: U74900KA2016PTC086221",
        "email": "info@tresconglobal.com",
        "web": "tresconglobal.com"
    },
    "mangalore": {
        "filename": "Trescon_Executive_Letterhead_Mangalore.docx",
        "blank_filename": "Trescon_Letterhead_Mangalore_Blank.docx",
        "entity": "Trescon Global Business Solutions Pvt Ltd.",
        "address": "1st Floor, Bejai Post, Ajantha Business Center, Bejai – Kapikad Road, Mangaluru, Karnataka – 575004",
        "extra": "",
        "cin": "CIN: U74900KA2016PTC086221",
        "email": "info@tresconglobal.com",
        "web": "tresconglobal.com"
    },
    "dubai": {
        "filename": "Trescon_Executive_Letterhead_Dubai.docx",
        "blank_filename": "Trescon_Letterhead_Dubai_Blank.docx",
        "entity": "Trescon Events Organizing Ltd.",
        "address": "Office 806, 8th Floor, Liberty House, Dubai International Financial Centre, DIFC, Dubai, UAE",
        "extra": "License number CL6668.",
        "cin": "",
        "email": "uae@tresconglobal.com",
        "web": "tresconglobal.com"
    }
}

def create_pixel_perfect_office_template(key, info, is_blank=False, output_filename=None):
    doc = docx.Document()
    
    # 1. Page Setup (Exact A4 Dimensions & 96 DPI Formula)
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    
    # Margins: Top 34mm (1927 dxa), Bottom 20mm (1134 dxa), Left/Right 10mm (567 dxa), Header 4mm (226 dxa)
    section.top_margin = Mm(34)      # 1927 dxa (Guarantees generous clearance below bottom rule)
    section.bottom_margin = Mm(20)   # 1134 dxa
    section.left_margin = Mm(10)     # 567 dxa
    section.right_margin = Mm(10)    # 567 dxa
    section.header_distance = Mm(4)  # 226 dxa (Positions top accent bar right at top edge)
    section.footer_distance = Mm(7)  # 397 dxa
    
    # Color Palette Definitions
    COLOR_TEAL = RGBColor(0, 165, 163)       # #00A5A3
    COLOR_DARK = RGBColor(1, 55, 61)        # #01373D
    COLOR_BODY = RGBColor(30, 33, 36)       # #1E2124
    COLOR_MUTED = RGBColor(100, 116, 139)   # #64748B
    
    # -------------------------------------------------------------
    # 2. HEADER SETUP (Clean 2-Column Layout Without Tagline)
    # -------------------------------------------------------------
    header = section.header
    
    # Top Accent Bar Paragraph at top edge of page
    p_top_bar = header.paragraphs[0]
    p_top_bar.paragraph_format.space_before = Pt(0)
    p_top_bar.paragraph_format.space_after = Pt(8)
    pBdr_top = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="36" w:space="0" w:color="01373D"/>'
        f'</w:pBdr>'
    )
    p_top_bar._p.get_or_add_pPr().append(pBdr_top)
    
    # Header Table (190mm Printable Width: Col 0 Logo 130mm, Col 1 Contact Meta 60mm)
    table = header.add_table(rows=1, cols=2, width=Mm(190))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    col_widths = [Mm(130), Mm(60)]
    for i, col in enumerate(table.columns):
        col.width = col_widths[i]
        for cell in col.cells:
            cell.width = col_widths[i]

    # Remove All Table Borders & Inner Gridlines
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:bottom w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="none"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(r'<w:tcBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>' % nsdecls('w'))
        tcPr.append(tcBorders)

    # Populate Header Content
    # Cell 0: Official Trescon Logo from OneDrive PNG Artboard 1 (Height = 60px @ 96DPI = 571,500 EMUs)
    cell_logo = table.cell(0, 0)
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_after = Pt(0)
    
    logo_path = 'brand_assets/10-years-trescon-logo.png'
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, height=Emu(571500))
    else:
        run_txt = p_logo.add_run("TRESCON GLOBAL")
        run_txt.font.name = 'Anek Devanagari'
        run_txt.font.size = Pt(14)
        run_txt.font.bold = True
        run_txt.font.color.rgb = COLOR_TEAL

    # Cell 1: Contact Meta Right Aligned (Email & Web)
    cell_meta = table.cell(0, 1)
    p_meta = cell_meta.paragraphs[0]
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_meta.paragraph_format.space_before = Pt(8)
    p_meta.paragraph_format.space_after = Pt(0)
    
    run_eicon = p_meta.add_run("✉ ")
    run_eicon.font.name = 'Segoe UI Symbol'
    run_eicon.font.size = Pt(9)
    run_eicon.font.color.rgb = COLOR_TEAL
    
    run_email = p_meta.add_run(info['email'] + "\n")
    run_email.font.name = 'Manrope'
    run_email.font.size = Pt(9)
    run_email.font.color.rgb = RGBColor(70, 77, 83)

    run_wicon = p_meta.add_run("🌐 ")
    run_wicon.font.name = 'Segoe UI Symbol'
    run_wicon.font.size = Pt(9)
    run_wicon.font.color.rgb = COLOR_TEAL
    
    run_web = p_meta.add_run(info['web'])
    run_web.font.name = 'Manrope'
    run_web.font.size = Pt(9)
    run_web.font.bold = True
    run_web.font.color.rgb = COLOR_TEAL

    # Teal Accent Rule Under Header
    p_rule = header.add_paragraph()
    p_rule.paragraph_format.space_before = Pt(8)
    p_rule.paragraph_format.space_after = Pt(0)
    p_rule_border = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="00A5A3"/>'
        f'</w:pBdr>'
    )
    p_rule._p.get_or_add_pPr().append(p_rule_border)

    # -------------------------------------------------------------
    # 3. FOOTER SETUP
    # -------------------------------------------------------------
    footer = section.footer
    
    footer_table = footer.add_table(rows=1, cols=1, width=Mm(190))
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    ftblPr = footer_table._tbl.tblPr
    ftblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="00A5A3"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/>'
        f'  <w:insideH w:val="none"/><w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    ftblPr.append(ftblBorders)
    
    cell_foot = footer_table.cell(0, 0)
    cell_foot.width = Mm(190)
    tcPr_f = cell_foot._tc.get_or_add_tcPr()
    tcBorders_f = parse_xml(r'<w:tcBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>' % nsdecls('w'))
    tcPr_f.append(tcBorders_f)
    
    p_f1 = cell_foot.paragraphs[0]
    p_f1.paragraph_format.space_before = Pt(4)
    p_f1.paragraph_format.space_after = Pt(2)
    run_ent = p_f1.add_run(info['entity'])
    run_ent.font.name = 'Anek Devanagari'
    run_ent.font.size = Pt(10.5)
    run_ent.font.bold = True
    run_ent.font.color.rgb = COLOR_DARK
    
    p_f2 = cell_foot.add_paragraph()
    p_f2.paragraph_format.space_before = Pt(0)
    p_f2.paragraph_format.space_after = Pt(2)
    run_a = p_f2.add_run(info['address'])
    run_a.font.name = 'Manrope'
    run_a.font.size = Pt(8.5)
    run_a.font.color.rgb = RGBColor(74, 85, 104)
    
    cin_extra_str = info['cin'] or info['extra']
    if cin_extra_str:
        p_f3 = cell_foot.add_paragraph()
        p_f3.paragraph_format.space_before = Pt(0)
        p_f3.paragraph_format.space_after = Pt(4)
        run_ce = p_f3.add_run(f"[{cin_extra_str}]")
        run_ce.font.name = 'Manrope'
        run_ce.font.size = Pt(8.0)
        run_ce.font.color.rgb = RGBColor(74, 85, 104)

    # Disclaimer Line
    p_disc = footer.add_paragraph()
    p_disc.paragraph_format.space_before = Pt(4)
    p_disc.paragraph_format.space_after = Pt(0)
    pBdr_disc = parse_xml(r'<w:pBdr %s><w:top w:val="single" w:sz="6" w:space="4" w:color="DCE3E6"/></w:pBdr>' % nsdecls('w'))
    p_disc._p.get_or_add_pPr().append(pBdr_disc)
    
    r_lbl = p_disc.add_run("Disclaimer: ")
    r_lbl.font.name = 'Manrope'
    r_lbl.font.size = Pt(7.5)
    r_lbl.font.bold = True
    r_lbl.font.color.rgb = COLOR_MUTED
    
    r_txt = p_disc.add_run(
        "The information shared by Trescon is confidential and intended solely for the recipient. "
        "It may not be copied, distributed, or relied upon without prior written consent. "
        "Trescon makes no warranties regarding the accuracy or completeness of the content and accepts no liability "
        "for any loss arising from its use. © 2025 Trescon. All rights reserved."
    )
    r_txt.font.name = 'Manrope'
    r_txt.font.size = Pt(7.5)
    r_txt.font.color.rgb = COLOR_MUTED

    # -------------------------------------------------------------
    # 4. BODY CONTENT
    # -------------------------------------------------------------
    if is_blank:
        p_b = doc.add_paragraph()
        p_b.paragraph_format.space_before = Pt(12)
        p_b.paragraph_format.space_after = Pt(0)
    else:
        p_recip = doc.add_paragraph()
        p_recip.paragraph_format.space_before = Pt(12)
        p_recip.paragraph_format.space_after = Pt(14)
        run_recip = p_recip.add_run(
            "To,\n"
            "Mr. Alex Turner\n"
            "Chief Executive Officer\n"
            "Apex Global Innovations Ltd.\n"
            "Bengaluru, Karnataka"
        )
        run_recip.font.name = 'Manrope'
        run_recip.font.size = Pt(10.5)
        run_recip.font.color.rgb = COLOR_BODY

        p_subj = doc.add_paragraph()
        p_subj.paragraph_format.space_after = Pt(12)
        p_subj.paragraph_format.left_indent = Pt(6)
        pBdr_subj = parse_xml(r'<w:pBdr %s><w:left w:val="single" w:sz="18" w:space="8" w:color="00A5A3"/></w:pBdr>' % nsdecls('w'))
        p_subj._p.get_or_add_pPr().append(pBdr_subj)
        
        run_subj = p_subj.add_run("Subject: Formal Proposal & Corporate Partnership Engagement")
        run_subj.font.name = 'Manrope'
        run_subj.font.size = Pt(10.5)
        run_subj.font.bold = True
        run_subj.font.color.rgb = COLOR_DARK

        p_sal = doc.add_paragraph()
        p_sal.paragraph_format.space_after = Pt(10)
        run_sal = p_sal.add_run("Dear Mr. Turner,")
        run_sal.font.name = 'Manrope'
        run_sal.font.size = Pt(10.5)
        run_sal.font.color.rgb = COLOR_BODY

        p_body1 = doc.add_paragraph()
        p_body1.paragraph_format.space_after = Pt(8)
        p_body1.paragraph_format.line_spacing = 1.15
        run_b1 = p_body1.add_run(
            "I hope this message finds you well. I am writing on behalf of Trescon Global to formally submit our "
            "comprehensive proposal for the upcoming enterprise technology summit and strategic collaboration initiatives. "
            "Our team has tailored this framework to align with your organization's vision, key deliverables, and expansion roadmaps."
        )
        run_b1.font.name = 'Manrope'
        run_b1.font.size = Pt(10.5)
        run_b1.font.color.rgb = COLOR_BODY

        p_body2 = doc.add_paragraph()
        p_body2.paragraph_format.space_after = Pt(8)
        p_body2.paragraph_format.line_spacing = 1.15
        run_b2 = p_body2.add_run(
            "As a premier B2B events and business solutions firm operating across seven global territories, Trescon is committed to "
            "connecting businesses with high-impact market opportunities. The enclosed document details our execution timeline, "
            "stakeholder engagement models, and target milestones for optimal business outcome."
        )
        run_b2.font.name = 'Manrope'
        run_b2.font.size = Pt(10.5)
        run_b2.font.color.rgb = COLOR_BODY

        p_sign = doc.add_paragraph()
        p_sign.paragraph_format.space_before = Pt(16)
        p_sign.paragraph_format.space_after = Pt(0)
        run_close = p_sign.add_run("Warm regards,\n\n\n")
        run_close.font.name = 'Manrope'
        run_close.font.size = Pt(10.5)
        run_close.font.color.rgb = COLOR_BODY

        run_name = p_sign.add_run("Mohammed Saleem\n")
        run_name.font.name = 'Manrope'
        run_name.font.size = Pt(10.5)
        run_name.font.bold = True
        run_name.font.color.rgb = COLOR_DARK

        run_title = p_sign.add_run("Founder & Chairman\nTrescon Global Business Solutions Pvt. Ltd.")
        run_title.font.name = 'Manrope'
        run_title.font.size = Pt(9.5)
        run_title.font.color.rgb = COLOR_MUTED

    fn = output_filename or (info['blank_filename'] if is_blank else info['filename'])
    doc.save(fn)
    print(f"Generated clean 2-column Word document with official OneDrive logo (no tagline): {fn}")

def main():
    for k, info in OFFICE_DATA.items():
        create_pixel_perfect_office_template(k, info, is_blank=False)
        create_pixel_perfect_office_template(k, info, is_blank=True)
    
    # Master Template
    create_pixel_perfect_office_template("bangalore", OFFICE_DATA["bangalore"], is_blank=True, output_filename="Trescon_Global_Letterhead_Template.docx")

if __name__ == '__main__':
    main()
